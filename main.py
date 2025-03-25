import re
from pathlib import Path
import argparse

from docx import Document


# \w matches all word characters
# \b - word boundaries
# + match repetitions
# flag re.ASCII to filter ASCII symbols
pattern = re.compile(r"\w+\b", flags=re.ASCII)


def collect_data() -> list:
    working_dir = Path("docx_files")
    resulted = []
    for f in working_dir.iterdir():
        if str(f).endswith("docx"):
            print(f"Processing {f}...")
            doc = Document(f)
            # collect all non-empty paragraphs with latin letters in doc
            _ = [
                resulted.append(p.text)
                for p in doc.paragraphs
                if len(p.text) > 0 and pattern.search(p.text)
            ]

    print(f"Collected {len(resulted)} paragraphs")
    return resulted


def clear_data() -> list:
    data = collect_data()

    data_to_str = []
    # convert list of lists to list of str and filer latin words
    for el in data:
        splitted = el.split(" ")
        _ = [data_to_str.append(s) for s in splitted if pattern.search(s)]
    # clean out non-breaking space in text \xc2 and \xa0
    non_braking_spaces = [
        s.replace("\xa0", " ").replace("\xc2", " ") if "\xa0" or "\xc2" in s else s
        for s in data_to_str
    ]
    # clean out arabic and latin digits from data
    resulted = [
        s
        for s in non_braking_spaces
        if not (bool(re.search(r"\d", s)) or s[:2].isupper())
    ]
    return resulted


def compile_names(ascii_str: list) -> list:
    """
    The function checks two first positions in the input list and
    returns combined species name. The logic is based on letter case checking.
    Species names always starts with capitalized word which is followed by lower case word.
    Pay attention that operations with the input list takes place not here
    but in the process() func.
    """
    if ascii_str[1].islower():
        res = ascii_str[0] + " " + ascii_str[1]
        i = 2
    elif ascii_str[1].startswith("("):
        res = ascii_str[0] + " " + ascii_str[1]
        i = 2
        if ascii_str[2].islower():
            res = res + " " + ascii_str[2]
        print(res, len(ascii_str))
    elif ascii_str[1][0].isupper():
        res = ascii_str[0]
        i = 1
    return (res, i)


def clear_result(data: list) -> set:
    cleared = []
    for s in data:
        if s.endswith("spp.,") or s.endswith("spp.."):
            s = s[:-1].strip()
            cleared.append(s)
        if ("spp" not in s) and (s.endswith(",") or s.endswith(".")):
            s = s[:-1].strip()
            cleared.append(s)
    # remove duplicate names
    cleared_set = sorted(set(cleared))
    print(cleared_set)
    return cleared_set


def result_to_file(names: set, filename: str):
    # using os.path.join looks shorter and clearer
    output_folder = Path("output")
    output_file = filename + ".docx"
    output_file_path = output_folder / output_file
    output_doc = Document()

    # add title
    title = output_doc.add_paragraph()
    t_run = title.add_run("List of Species")
    t_run.bold = True

    # compile list of species
    for name in names:
        output_doc.add_paragraph(name)

    output_doc.save(output_file_path)
    print(f"\n List of species is saved to {output_file_path}")


def process(filename: str):
    comb_names = []
    ascii_str = clear_data()

    while len(ascii_str) > 1:
        s, i = compile_names(ascii_str)
        comb_names.append(s)
        del ascii_str[:i]
    else:
        if ascii_str:
            if ascii_str[0].isupper() or ascii_str[0].startswith("("):
                comb_names.append(ascii_str[0])
            del ascii_str[0]
        print(sorted(comb_names))
        print("Task completed.")

    resulted = clear_result(comb_names)
    result_to_file(resulted, filename)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "-f",
        "--filename",
        type=str,
        help="""Name of output file containing latin species names (without file extension)""",
    )
    args = parser.parse_args()

    if args.filename:
        filename = args.filename
    else:
        filename = "species_list"

    process(filename)
